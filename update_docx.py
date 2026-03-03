import sys
from docx import Document
from docx.oxml.ns import qn
from copy import deepcopy

doc = Document('content.docx')

# ── Configuration: define quiz content here ──────────────────────────────
QUIZ_DATA = {
    'question': 'Which of the following was NOT a factor in the making of a global world?',
    'options': [
        'A) Silk Routes',
        'B) Discovery of the Americas',
        'C) Industrial Revolution',
        'D) Invention of the Internet'
    ]
}

# ── Find the [sst_quiztime_page] section and inject content ─────────────
target_section = '[sst_quiztime_page]'
found = False

for i, para in enumerate(doc.paragraphs):
    if para.text.strip() == target_section:
        found = True
        # Find the next section header or end of document
        insert_after = para

        # Remove any existing content between this header and the next section
        j = i + 1
        while j < len(doc.paragraphs):
            next_text = doc.paragraphs[j].text.strip()
            if next_text.startswith('[') and next_text.endswith(']'):
                break  # Hit next section
            # Clear existing content
            doc.paragraphs[j].text = ''
            j += 1

        # Insert quiz question and options after the header
        # We'll reuse the cleared paragraphs or add text to them
        content_lines = [f"Question: {QUIZ_DATA['question']}"]
        for opt in QUIZ_DATA['options']:
            content_lines.append(opt)

        # Fill cleared paragraphs or use the first available one
        next_idx = i + 1
        for line_idx, line in enumerate(content_lines):
            if next_idx + line_idx < len(doc.paragraphs):
                target_para = doc.paragraphs[next_idx + line_idx]
                if not target_para.text.strip().startswith('['):
                    target_para.text = line
                    continue
            # If we've run out of cleared paragraphs, add new ones after the header
            new_p = deepcopy(insert_after._element)
            new_p.clear()
            # Add a run with the text
            r = new_p.makeelement(qn('w:r'), {})
            t = r.makeelement(qn('w:t'), {})
            t.text = line
            r.append(t)
            new_p.append(r)
            insert_after._element.addnext(new_p)
            insert_after = type(insert_after)(new_p, insert_after._parent)

        print(f"Updated {target_section} with:")
        print(f"  Question: {QUIZ_DATA['question']}")
        print(f"  Options:  {', '.join(QUIZ_DATA['options'])}")
        break

if not found:
    print(f"ERROR: Section {target_section} not found in content.docx")
    sys.exit(1)

try:
    doc.save('content.docx')
    print("Saved content.docx successfully.")
except PermissionError:
    doc.save('content_updated.docx')
    print("content.docx is locked (open in Word?). Saved as content_updated.docx instead.")
