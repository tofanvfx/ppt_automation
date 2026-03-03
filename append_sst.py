from docx import Document

doc = Document('content.docx')
doc.add_paragraph("[sst_title_page]")
doc.add_paragraph("CLASS: Class - 10")
doc.add_paragraph("SUBJECT: Social Science")
doc.add_paragraph("CHAPTER_NUMBER: Chapter - 3")
doc.add_paragraph("CHAPTER_NAME: Making of a Global World")
doc.add_paragraph("LESSON: Lesson - 1")
doc.add_paragraph("TOPIC: The Pre-Modern World")
doc.save('content.docx')
print("Added [sst_title_page] to content.docx")
